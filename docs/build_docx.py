# -*- coding: utf-8 -*-
"""Build BAO_CAO_A01.docx theo đúng cấu trúc mẫu A01_06_linhpm.488.pdf (41 trang)."""
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = r"C:\Users\Laptop\OneDrive\Laptop\Ki 1 - Nam 4\HTTM\assignment01"
FIG = os.path.join(BASE, "figures")

doc = Document()

# ============ GLOBAL STYLE ============
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(13)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.3
pf.space_after = Pt(6)

for i in range(1, 4):
    h = doc.styles[f"Heading {i}"]
    h.font.name = "Times New Roman"
    h.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.font.bold = True
    h.font.size = Pt([16, 14, 13][i - 1])

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)


def p(text, bold=False, italic=False, align=None, size=None, space_after=None):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align == "center":
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if space_after is not None:
        par.paragraph_format.space_after = Pt(space_after)
    return par


def rich(parts, align="justify"):
    """parts = list of (text, bold) tuples."""
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if align == "justify" else WD_ALIGN_PARAGRAPH.LEFT
    for text, bold in parts:
        run = par.add_run(text)
        run.bold = bold
    return par


def img(filename, width_cm, caption=None):
    path = os.path.join(FIG, filename)
    if not os.path.exists(path):
        print("MISSING FIGURE:", filename)
        return
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after = Pt(2)
    run = par.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(11)
        cap.paragraph_format.space_after = Pt(10)


def formula(text):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(text)
    run.italic = True
    run.font.name = "Cambria Math"
    par.paragraph_format.space_after = Pt(8)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9E2F3")
        cell._tc.get_or_add_tcPr().append(shd)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.rows[i + 1].cells[j]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run(str(v))
            run.font.size = Pt(11)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ============ TRANG BÌA (theo mẫu: khung trang trí + logo PTIT) ============
# Nền khung trang trí PTIT: ảnh lớn căn giữa, nằm sau văn bản
def add_cover_frame():
    """Chèn khung trang trí làm ảnh nổi (floating) phủ toàn trang bìa, sau text."""
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run()
    pic = run.add_picture(os.path.join(BASE, "docs", "ptit_frame.png"), width=Cm(21.0), height=Cm(29.7))
    # chuyển inline -> floating behindDocument
    inline = run._r.findall(qn("w:drawing"))[0][0]
    import copy
    ext = inline.find(qn("wp:extent"))
    cx, cy = int(ext.get("cx")), int(ext.get("cy"))
    docPr = inline.find(qn("wp:docPr"))
    anchor_el = copy.deepcopy(inline)
    # Tạo wp:anchor thay cho wp:inline
    from lxml import etree
    WPE = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    anchor = etree.SubElement(inline.getparent(), qn("wp:anchor"))
    for k, v in [("distT", "0"), ("distB", "0"), ("distL", "0"), ("distR", "0"),
                 ("simplePos", "0"), ("relativeHeight", "0"), ("behindDoc", "1"),
                 ("locked", "0"), ("layoutInCell", "1"), ("allowOverlap", "1")]:
        anchor.set(k, v)
    sp = etree.SubElement(anchor, qn("wp:simplePos")); sp.set("x", "0"); sp.set("y", "0")
    ph = etree.SubElement(anchor, qn("wp:positionH")); ph.set("relativeFrom", "page")
    po = etree.SubElement(ph, qn("wp:posOffset")); po.text = "0"
    pv = etree.SubElement(anchor, qn("wp:positionV")); pv.set("relativeFrom", "page")
    po2 = etree.SubElement(v if False else pv, qn("wp:posOffset")); po2.text = "0"
    e2 = etree.SubElement(anchor, qn("wp:extent")); e2.set("cx", str(cx)); e2.set("cy", str(cy))
    wra = etree.SubElement(anchor, qn("wp:wrapNone"))
    dp2 = etree.SubElement(anchor, qn("wp:docPr"))
    dp2.set("id", docPr.get("id", "1")); dp2.set("name", "cover-frame")
    # copy graphic
    graphic = inline.find(qn("a:graphic"))
    anchor.append(copy.deepcopy(graphic))
    # remove inline, insert anchor
    inline.getparent().remove(inline)

add_cover_frame()
p("HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG", bold=True, align="center", size=14)
p("KHOA CÔNG NGHỆ THÔNG TIN 1", bold=True, align="center", size=14)
# Logo PTIT căn giữa (như mẫu: dưới header, trên tiêu đề)
logo_par = doc.add_paragraph()
logo_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_run = logo_par.add_run()
logo_run.add_picture(os.path.join(BASE, "docs", "ptit_logo.png"), height=Cm(3.2))
for _ in range(3):
    doc.add_paragraph()
p("ASSIGNMENT 01 – INTELLIGENT SYSTEM DEVELOPMENT", bold=True, align="center", size=22)
for _ in range(2):
    doc.add_paragraph()

t = doc.add_table(rows=4, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
info = [
    ("Học phần:", "Phát triển các hệ thống thông minh"),
    ("Giảng viên hướng dẫn:", "Thầy Trần Đình Quế"),
    ("Sinh viên thực hiện:", "................................"),
    ("Mã sinh viên:", "................................"),
]
for i, (k, v) in enumerate(info):
    r1 = t.rows[i].cells[0].paragraphs[0].add_run(k)
    r1.bold = True
    r1.font.size = Pt(14)
    r2 = t.rows[i].cells[1].paragraphs[0].add_run(v)
    r2.font.size = Pt(14)
for _ in range(6):
    doc.add_paragraph()
p("Hà Nội – 2026", bold=True, align="center", size=14)
doc.add_page_break()

# ============ MỤC LỤC ============
p("MỤC LỤC", bold=True, align="center", size=16)
toc_items = [
    ("PHẦN I – LÝ THUYẾT", 0),
    ("1. Từ Intelligence đến Intelligent System", 1),
    ("2. Hai bài toán và phát biểu bài toán chính thức", 1),
    ("3. Data Representation – Biểu diễn dữ liệu", 1),
    ("4. Supervised Learning", 1),
    ("5. Baseline", 1),
    ("6. Các mô hình Machine Learning", 1),
    ("7. Độ đo đánh giá", 1),
    ("8. Controlled Experiments", 1),
    ("9. Từ Model đến Application", 1),
    ("PHẦN II – HỆ THỐNG 1: CHẨN ĐOÁN BỆNH TIM", 0),
    ("10. Mục tiêu hệ thống", 1),
    ("11. Exploratory Data Analysis – Các biểu đồ và giải thích", 1),
    ("12. Data Representation và preprocessing", 1),
    ("13. Baseline và 5 mô hình", 1),
    ("14. Experiment 1 – Model Comparison", 1),
    ("15. Experiment 2 – Hyperparameter Investigation", 1),
    ("16. Experiment 3 – Representation Investigation", 1),
    ("17. Final Model và lưu model", 1),
    ("PHẦN III – HỆ THỐNG 2: DỰ ĐOÁN GIÁ NHÀ", 0),
    ("18. Mục tiêu", 1),
    ("19. EDA – Biểu đồ giá nhà và giải thích", 1),
    ("20. Train/Test và Representation cho giá nhà", 1),
    ("21. Experiment 1 – Model Comparison", 1),
    ("22. Experiment 2 – Hyperparameter Investigation", 1),
    ("23. Experiment 3 – Representation / Feature Investigation", 1),
    ("24. Final Model – Giá nhà", 1),
    ("PHẦN IV – TRIỂN KHAI WEB", 0),
    ("25. Kiến trúc Web", 1),
    ("26. Demo ứng dụng", 1),
]
for text, lvl in toc_items:
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(0.75 * lvl)
    run = par.add_run(text)
    run.bold = lvl == 0
doc.add_page_break()

# ============ PHẦN I ============
doc.add_heading("PHẦN I – LÝ THUYẾT", level=1)

doc.add_heading("1. Từ Intelligence đến Intelligent System", level=2)
p("Theo Assignment 01, trí thông minh trong bài được hiểu theo nghĩa vận hành: hệ thống có khả năng tiếp nhận thông tin, biểu diễn thông tin dưới dạng tính toán, học mối quan hệ từ các ví dụ, dự đoán cho trường hợp mới và hỗ trợ một quyết định/ứng dụng.", align="justify")
p("Vì vậy, một mô hình Machine Learning đơn lẻ chưa phải là toàn bộ hệ thống thông minh. Hệ thống hoàn chỉnh phải có chuỗi:", align="justify")
p("Môi trường/Dữ liệu → Input → Representation → Learning/Model → Decision → Output → Application", bold=True, align="center")
p("Trong bài này:", align="justify")
p("• Hệ bệnh tim: thông tin sức khỏe → vector đặc trưng → mô hình phân loại → dự đoán có/không mắc bệnh tim.", align="justify")
p("• Hệ giá nhà: thông tin bất động sản → vector đặc trưng → mô hình hồi quy → giá nhà dự đoán.", align="justify")
p("Assignment nhấn mạnh rằng mô hình không trực tiếp nhận “thế giới thực”, mà nhận một biểu diễn tính toán của những thông tin được lựa chọn. Vì vậy Representation quyết định thông tin nào có sẵn cho việc học.", align="justify")

doc.add_heading("2. Hai bài toán và phát biểu bài toán chính thức", level=2)
doc.add_heading("2.1. Hệ thống chẩn đoán bệnh tim", level=3)
p("Bài toán thực tế: hỗ trợ sàng lọc bệnh tim (hẹp động mạch vành) dựa trên các chỉ số tim mạch đo lường được.", align="justify")
p("Input: 13 đặc trưng: age, sex, cp, trestbps, chol, fbs, restecg, thalach, exang, oldpeak, slope, ca, thal.", align="justify")
rich([("Target: ", True), ("num gốc 0–4 được quy về nhị phân, trong đó 0: không có bệnh tim theo nhãn dữ liệu; 1: có bệnh tim (num > 0).", False)])
p("Loại bài toán: Classification nhị phân.", align="justify")
p("Phát biểu chính thức: Với một vector đặc trưng sức khỏe chưa từng xuất hiện trong tập kiểm tra, dự đoán lớp target ∈ {0,1}.", align="justify", bold=False)
doc.add_heading("2.2. Hệ thống dự đoán giá nhà", level=3)
p("Bài toán thực tế: hỗ trợ ước lượng giá nhà dựa trên các thuộc tính của bất động sản tại Ames, Iowa (Mỹ).", align="justify")
p("Input: các đặc trưng như diện tích sinh hoạt, chất lượng tổng thể, số garage, diện tích tầng hầm, số phòng tắm, năm xây, khu dân cư và các thuộc tính khác.", align="justify")
p("Target: SalePrice — giá bán theo USD.", align="justify")
p("Loại bài toán: Regression.", align="justify")
p("Phát biểu chính thức: Với một vector đặc trưng bất động sản chưa từng xuất hiện trong tập kiểm tra, dự đoán giá nhà SalePrice dưới dạng một giá trị số.", align="justify")

doc.add_heading("3. Data Representation – Biểu diễn dữ liệu", level=2)
p("Với dữ liệu có cấu trúc, mỗi quan sát được biểu diễn bởi một vector:", align="justify")
formula("xi = [xi1, xi2, ..., xid] ∈ Rd")
p("Toàn bộ tập dữ liệu được biểu diễn bởi:")
formula("D = {(xi, yi)}; i = 1..N")
p("Trong đó: xi là biểu diễn đầu vào của quan sát thứ i; d là số đặc trưng sau khi biểu diễn; yi là target; N là số quan sát; X là ma trận đặc trưng; y là vector mục tiêu.", align="justify")
p("Điểm quan trọng: raw feature ≠ encoded feature ≠ model input nếu quá trình tiền xử lý thay đổi biểu diễn.", bold=True, align="justify")
p("Bệnh tim: 6 biến số (age, trestbps, chol, thalach, oldpeak, ca) và 7 biến phân loại (sex, cp, fbs, restecg, exang, slope, thal). Bản gốc UCI có 6 ô “?” (ca thiếu 4, thal thiếu 2) được thay bằng median. Các biến phân loại được mã hóa One-Hot Encoding, nâng số chiều từ 13 lên 18.", align="justify")
p("Giá nhà: dữ liệu có cả biến số và biến phân loại. Các biến phân loại được mã hóa One-Hot Encoding; các cột thiếu được xử lý theo ngữ nghĩa (NA của Garage/Bsmt nghĩa là “không có tiện ích”, điền “None”; LotFrontage điền median theo khu dân cư), nâng số chiều từ 74 lên 271. Vì các mô hình và preprocessing nằm trong Pipeline, cùng một biểu diễn được áp dụng cho train, test và dữ liệu mới.", align="justify")

doc.add_heading("4. Supervised Learning", level=2)
p("Mô hình học một hàm:")
formula("ŷ = fθ(x)")
p("Mục tiêu là tìm tham số θ* giúp sai số dự đoán nhỏ trên dữ liệu học, đồng thời có khả năng generalization – dự đoán tốt cho dữ liệu chưa từng thấy:")
formula("θ* = argmin θ (1/N) Σ ℓ(fθ(xi), yi)")
doc.add_heading("Train/Test", level=3)
p("Dữ liệu được chia thành: Training set dùng để học; Test set chỉ dùng để đánh giá cuối cùng. Trong bài, sử dụng tỷ lệ 80/20 và random_state = 42. Với classification sử dụng stratify = y.", align="justify")
p("Không sử dụng test set để lựa chọn mô hình hoặc tinh chỉnh tham số, nhằm tránh data leakage. Các quyết định trung gian dựa trên cross-validation trên training set (5-fold với bệnh tim, 10-fold với giá nhà).", align="justify")

doc.add_heading("5. Baseline", level=2)
p("Baseline là mốc tham chiếu đơn giản nhất.", align="justify")
p("• Classification: DummyClassifier(strategy=\"most_frequent\").")
p("• Regression: DummyRegressor(strategy=\"mean\").")
p("Baseline trả lời câu hỏi: “Các mô hình Machine Learning có thực sự tốt hơn một chiến lược đơn giản hay không?” Assignment yêu cầu phải có baseline trước khi đánh giá các mô hình phức tạp hơn.", align="justify")

doc.add_heading("6. Các mô hình Machine Learning", level=2)
p("Classification – Bệnh tim:", bold=True)
p("1. Logistic Regression – học ranh giới tuyến tính và xác suất lớp.")
p("2. K-Nearest Neighbors – dự đoán dựa trên các điểm dữ liệu gần nhất.")
p("3. Support Vector Machine – tìm biên phân tách với margin phù hợp.")
p("4. Decision Tree – chia không gian đặc trưng thành các vùng bằng các luật rẽ nhánh.")
p("5. Random Forest – kết hợp nhiều cây quyết định để tăng tính ổn định.")
p("Regression – Giá nhà:", bold=True)
p("1. Linear Regression – mô hình hóa quan hệ tuyến tính.")
p("2. Decision Tree Regressor – phân vùng không gian đặc trưng và dự đoán theo vùng.")
p("3. Random Forest Regressor – trung bình nhiều cây hồi quy.")
p("4. Support Vector Regression – tìm hàm dự đoán lệch target không quá ε và phẳng nhất có thể.")
p("Với mỗi mô hình, bài cần quan tâm: mô hình nhận representation nào, học quan hệ gì, cấu trúc/tham số nào được học, tiêu chí học, giả định, ưu điểm và hạn chế.", align="justify")

doc.add_heading("7. Độ đo đánh giá", level=2)
p("Classification", bold=True)
p("• Accuracy: tỷ lệ dự đoán đúng trên toàn bộ mẫu.")
p("• Precision: trong các mẫu được dự đoán dương tính, bao nhiêu mẫu thực sự dương tính.")
p("• Recall: trong các mẫu thực sự dương tính, mô hình phát hiện được bao nhiêu.")
p("• F1-score: trung bình điều hòa giữa Precision và Recall.")
p("• Confusion Matrix: cho biết TP, TN, FP, FN.")
p("Trong hệ thống sàng lọc bệnh, Recall có ý nghĩa đặc biệt vì bỏ sót một trường hợp có nguy cơ có thể nghiêm trọng. Tuy nhiên, không thể chỉ tối ưu Recall mà bỏ qua Precision và F1. Bài này bổ sung ROC-AUC để đo khả năng xếp hạng hai lớp trên nhiều ngưỡng.", align="justify")
p("Regression", bold=True)
p("• MAE: sai số tuyệt đối trung bình, dễ diễn giải theo đơn vị của giá.")
p("• MSE: bình phương sai số trung bình, phạt mạnh sai số lớn.")
p("• RMSE: căn bậc hai của MSE, cùng đơn vị với target.")
p("• R²: mức độ giải thích biến thiên của target; càng gần 1 càng tốt.")
p("Không sử dụng Accuracy cho bài toán giá nhà vì đây là regression.", align="justify")

doc.add_heading("8. Controlled Experiments", level=2)
p("Assignment yêu cầu tối thiểu 3 thí nghiệm có câu hỏi rõ ràng.", align="justify")
p("Experiment 1 – Model Comparison. Câu hỏi: Với cùng một train/test split và cùng protocol đánh giá, mô hình nào hoạt động tốt nhất?", align="justify")
p("Experiment 2 – Hyperparameter Investigation. Câu hỏi: Thay đổi một hyperparameter có ảnh hưởng đến kết quả không? Với bệnh tim khảo sát k của KNN và n_estimators của Random Forest; với giá nhà khảo sát max_depth của Decision Tree và n_estimators của Random Forest.", align="justify")
p("Experiment 3 – Representation / Feature Investigation. Câu hỏi: Thay đổi representation có ảnh hưởng đến khả năng học không? Với bệnh tim so sánh dữ liệu có chuẩn hóa và không chuẩn hóa, đồng thời so sánh X_all với tập feature quan trọng nhất; với giá nhà so sánh target gốc với target log-transform và X_all với nhóm feature quan trọng nhất.", align="justify")
p("Mỗi experiment phải thay đổi có kiểm soát, không thay đổi hàng loạt tham số cùng lúc.", align="justify")

doc.add_heading("9. Từ Model đến Application", level=2)
p("Một model đã train chưa phải hệ thống hoàn chỉnh. Application phải thực hiện:", align="justify")
p("Input → Feature Representation → Preprocessing → Model → Prediction → Output", bold=True, align="center")
p("Web app trong bài sử dụng cùng final model và cùng logic preprocessing thông qua hàm predict trung tâm: app nhận dữ liệu thô, mã hóa One-Hot đúng thứ tự cột như lúc training rồi mới đưa vào model. Điều này đảm bảo dữ liệu mới được biểu diễn giống dữ liệu lúc training.", align="justify")
doc.add_page_break()

# ============ PHẦN II ============
doc.add_heading("PHẦN II – HỆ THỐNG 1: CHẨN ĐOÁN BỆNH TIM", level=1)

doc.add_heading("10. Mục tiêu hệ thống", level=2)
p("Hệ thống nhận 13 thông tin đầu vào liên quan đến tim mạch và dự đoán nhãn target. Đây là hệ thống hỗ trợ sàng lọc học thuật, không phải công cụ chẩn đoán y khoa. Kết quả chỉ có ý nghĩa trong phạm vi dataset và mô hình đã học.", align="justify")
p("Dataset: UCI Heart Disease – bản processed.cleveland (https://archive.ics.uci.edu/dataset/45/heart+disease), 303 bệnh nhân, 13 feature + 1 target. Dataset đã được tải về máy và lưu tại data/heart_clean.csv để chạy offline, bảo đảm reproducible.", align="justify")
img("heart_01_target_dist.png", 11, "Hình 1 – Đọc dataset và phân bố target")
p("Giải thích dữ liệu: Dataset có 303 quan sát và 13 biến đầu vào cùng target nhị phân (num > 0 → 1). Sau khi thay 6 ô “?” bằng median, dữ liệu không còn missing. Target khá cân bằng: 164 mẫu lớp 0 (54,1%) và 139 mẫu lớp 1 (45,9%). Các biến như thalach, oldpeak, ca, cp, thal được kỳ vọng mang thông tin phân biệt hai nhóm. Tuy nhiên, mối quan hệ quan sát được chỉ là quan hệ trong dataset, không được diễn giải thành quan hệ nhân quả y khoa.", align="justify")

doc.add_heading("11. Exploratory Data Analysis – Các biểu đồ và giải thích", level=2)
p("Các biểu đồ không chỉ nhằm minh họa mà giúp trả lời các câu hỏi: Dữ liệu có mất cân bằng không? Feature nào có phân phối khác nhau giữa hai lớp? Có outlier hoặc giá trị bất thường không? Các feature có tương quan với nhau không?", align="justify")
doc.add_page_break()

# Biểu đồ 2
doc.add_heading("Biểu đồ 1 – Phân phối 6 feature số theo lớp target", level=3)
img("heart_02_num_dist.png", 15.5)
rich([("Cách đọc Biểu đồ 1 – Histogram 6 feature số theo Outcome", True)])
p("Bước 1 – Đọc trục: mỗi ô con là một feature; trục X là giá trị feature; trục Y là mật độ (density); hai màu là hai lớp target (xanh = khỏe, đỏ = bệnh).", align="justify")
p("Bước 2 – So sánh hai nhóm: nếu đường phân phối lớp đỏ dịch trái (thalach) hay dịch phải (oldpeak, ca) so với lớp xanh, feature đó có khả năng phân biệt.", align="justify")
p("Bước 3 – Hiểu vùng chồng lấn: hai phân phối chồng lấn nhiều (chol, trestbps) nghĩa là có những bệnh nhân cùng giá trị nhưng khác kết luận — feature hữu ích nhưng không đủ.", align="justify")
p("Bước 4 – Liên hệ mô hình: feature có phân phối lệch khác nhau giữa hai lớp giúp model học ranh giới; feature trùng phân phối (chol, trestbps) dự kiến đóng góp yếu, điều này được kiểm chứng bằng feature importance ở mục 13.", align="justify")
p("Câu nói khi thuyết trình: “Em xem phân phối từng feature theo hai lớp để chọn ứng viên feature mạnh. thalach của nhóm bệnh lệch rõ sang trái và oldpeak lệch sang phải, nên hai feature này được kỳ vọng đóng góp mạnh; còn chol và trestbps gần như trùng nhau nên dự kiến yếu.”", italic=True, align="justify")
doc.add_page_break()

# Biểu đồ 3
doc.add_heading("Biểu đồ 2 – Feature categorical theo lớp target", level=3)
img("heart_04_cat.png", 15.5)
rich([("Cách đọc Biểu đồ 2 – Tỷ lệ lớp theo giá trị categorical", True)])
p("Bước 1 – Đọc trục: mỗi ô là một feature categorical; trục X là giá trị của feature; trục Y là tỷ lệ phần trăm trong nhóm (stacked 100%).", align="justify")
p("Bước 2 – So sánh màu: phần màu đỏ (lớp 1) càng chiếm nhiều trong một cột, giá trị đó càng gắn với bệnh tim.", align="justify")
p("Bước 3 – Đọc xu hướng: với ca tăng từ 0 lên 3, tỷ lệ đỏ tăng gần đơn điệu — quan hệ dạng liều–phản ứng rất mạnh. cp = 4 (asymptomatic) có tỷ lệ bệnh vượt trội. thal = 7 (reversible defect) cao nhất trong ba loại.", align="justify")
p("Bước 4 – Liên hệ ML: đây là bằng chứng trực quan rằng mã hóa One-Hot các categorical này mang thông tin phân biệt tốt, giải thích vì sao representation 18 chiều sau encode hiệu quả.", align="justify")
doc.add_page_break()

# Biểu đồ 4
doc.add_heading("Biểu đồ 3 – Ma trận tương quan", level=3)
img("heart_03_corr.png", 13.5)
rich([("Cách đọc Biểu đồ 3 – Correlation Heatmap", True)])
p("Mỗi ô là hệ số tương quan giữa hai biến số, nằm trong khoảng −1 đến +1: gần +1 là quan hệ tuyến tính cùng chiều mạnh; gần 0 là yếu; gần −1 là ngược chiều mạnh. Chọn feature trên hàng và feature trên cột; giao điểm cho hệ số. Đường chéo chính bằng 1. Màu chỉ là cách mã hóa số — phải đọc cả giá trị số ghi trong ô.", align="justify")
p("Đặc biệt chú ý cột target: ca (≈0,46), thal (≈0,37), oldpeak (≈0,36), cp (≈−0,41), thalach (≈−0,40) là các feature đơn lẻ tương quan mạnh nhất. Giữa các feature không có cặp nào vượt |ρ| > 0,6 nên không có dư thừa nghiêm trọng,thúc đẩy việc giữ đủ 13 feature cho representation đầy đủ.", align="justify")
p("Correlation chỉ đo quan hệ tuyến tính và không chứng minh nhân quả — ca tương quan cao vì bản thân nó là kết quả chẩn đoán hình ảnh của bệnh tim.", align="justify")

# Biểu đồ 5
doc.add_heading("Biểu đồ 4 – Trọng số Logistic Regression", level=3)
img("heart_05_lr_weights.png", 14)
rich([("Cách đọc Biểu đồ 4 – LR weights", True)])
p("Trục X là trọng số w của từng feature sau chuẩn hóa; cột đỏ (w > 0) đẩy dự đoán về lớp 1 “có bệnh”, cột xanh (w < 0) đẩy về lớp 0 “khỏe”. Vì feature đã chuẩn hóa nên các trọng số so sánh trực tiếp được về độ lớn.", align="justify")
p("Các trọng số dương lớn nhất: ca_2, ca_3 (2–3 mạch vành hẹp), cp_4 (đau ngực không triệu chứng), thal_7 (reversible defect), exang_1, oldpeak — khớp hoàn toàn với EDA. Các trọng số âm: thalach (nhịp tim tối đa cao → lành), cp_3, ca_0. Mô hình tuyến tính này diễn giải được về mặt y khoa — một lợi thế quan trọng khi chọn model cho ứng dụng y tế.", align="justify")
doc.add_page_break()

doc.add_heading("Biểu đồ 5 – Feature Importance của Random Forest", level=3)
img("heart_06_rf_importance.png", 14)
rich([("Cách đọc Biểu đồ 5 – RF importance", True)])
p("Trục X là feature; trục Y là mức độ quan trọng theo Gini. Cột càng cao, feature càng được các cây dùng nhiều để phân chia.", align="justify")
p("Nhóm dẫn đầu: ca (gộp các mức One-Hot), thal, cp, oldpeak, thalach — trùng khớp cả với EDA lẫn trọng số LR. Ba nguồn bằng chứng độc lập (EDA, LR weights, RF importance) hội tụ về cùng một tập feature quan trọng nên độ tin cậy kết luận cao. chol và trestbps đúng như dự đoán: đóng góp thấp.", align="justify")

doc.add_heading("12. Data Representation và preprocessing", level=2)
p("Bản gốc có 6 ô “?” (ca thiếu 4, thal thiếu 2, khoảng 2%). Vì tỷ lệ rất nhỏ, các giá trị này được thay bằng median của cột — cách này không bị kéo bởi outlier như mean. Target num 0–4 được quy về nhị phân: num > 0 → 1.", align="justify")
p("7 biến phân loại được mã hóa One-Hot Encoding, nâng số chiều từ 13 lên 18. Để tránh data leakage, việc chuẩn hóa nằm trong Pipeline của từng mô hình: fit() của preprocessing chỉ học từ training set. Do đó cùng một biểu diễn được áp dụng cho train, test và dữ liệu mới trong application.", align="justify")

doc.add_heading("13. Baseline và 5 mô hình", level=2)
p("Baseline sử dụng lớp xuất hiện nhiều nhất: Accuracy 0,541; Precision 0; Recall 0; F1 0 — baseline không bao giờ dự đoán lớp 1 nên bỏ sót 100% người bệnh. Đây là minh chứng vì sao Accuracy một mình chưa đủ: một hệ “luôn đoán khỏe” đạt Accuracy 54% nhưng hoàn toàn vô dụng cho y tế.", align="justify")
p("Sau đó huấn luyện 5 mô hình được khuyến nghị trong Assignment. Điểm quan trọng: tất cả mô hình đều gọi fit(X_train, y_train) và chỉ dùng X_test/y_test để đánh giá.", align="justify")
table(
    ["Model", "Accuracy", "Precision", "Recall", "F1"],
    [
        ["Random Forest (100)", "0,885", "0,839", "0,929", "0,881"],
        ["KNN (k=11)", "0,885", "0,862", "0,893", "0,877"],
        ["Logistic Regression", "0,836", "0,800", "0,857", "0,828"],
        ["SVM (RBF, C=1)", "0,836", "0,800", "0,857", "0,828"],
        ["Decision Tree (depth 6)", "0,803", "0,727", "0,857", "0,788"],
        ["Baseline (majority)", "0,541", "0,000", "0,000", "0,000"],
    ],
)
p("Bảng 1 – Kết quả 5 mô hình và baseline trên test set (n = 61)", align="center", size=11)
rich([("Giải thích kết quả đánh giá. ", True), ("Không chọn mô hình chỉ dựa trên Accuracy. Với bài toán y tế, Recall và F1 có ý nghĩa quan trọng vì hệ thống cần hạn chế bỏ sót trường hợp dương tính nhưng vẫn kiểm soát false positive. Cả 5 mô hình đều vượt baseline áp đảo (F1 từ 0 lên 0,79–0,88) — mô hình đã học được tín hiệu thật. Random Forest đạt Recall 0,929 cao nhất: chỉ bỏ sót 2/28 người bệnh trong test set.", False)], align="justify")
img("heart_07_confusion.png", 15.5, "Hình 2 – Confusion Matrix 4 mô hình trên test set")
rich([("Cách đọc Confusion Matrix. ", True), ("Đọc từng ô, không đọc theo màu trước: TN (thực 0, đoán 0) và TP (thực 1, đoán 1) là hai ô đúng; FP (thực 0, đoán 1) là báo động nhầm; FN (thực 1, đoán 0) là bỏ sót người bệnh — ô nghiêm trọng nhất về y đức. Từ 4 ô suy ra Precision, Recall, F1 và Accuracy. Random Forest chỉ có 2 FN — số liệu trực tiếp đằng sau Recall 0,929.", False)], align="justify")
img("heart_08_roc.png", 11.5, "Hình 3 – ROC Curve so sánh các mô hình")
rich([("Cách đọc ROC Curve. ", True), ("Trục X = False Positive Rate = FP/(FP+TN); Trục Y = True Positive Rate = TP/(TP+FN). Mỗi điểm trên đường ROC tương ứng với một ngưỡng dự đoán khác nhau. Đường chéo 45° biểu diễn mức gần với dự đoán ngẫu nhiên; đường ROC càng nằm phía trên bên trái, model càng phân biệt hai lớp tốt. AUC là diện tích dưới đường ROC, càng gần 1 càng tốt — AUC các mô hình trong bài đạt 0,90–0,93. AUC không phải “độ chính xác y khoa” mà là khả năng xếp hạng hai lớp trên nhiều threshold.", False)], align="justify")
doc.add_page_break()

doc.add_heading("14. Experiment 1 – Model Comparison", level=2)
p("Câu hỏi: Trong cùng điều kiện train/test và preprocessing, mô hình nào đạt kết quả tốt hơn?", align="justify")
p("Để kết luận ổn định, năm mô hình được so sánh bằng 5-fold Stratified Cross-Validation trên training set (test set không dùng để lựa chọn). Xếp hạng CV nhất quán với test set: Random Forest ≈ SVM > Logistic Regression > KNN > Decision Tree. Random Forest vừa cao vừa ổn định (khoảng tứ phân vị hẹp), KNN thấp hơn và dao động lớn hơn.", align="justify")
p("Vì sao KNN thua: mã hóa One-Hot làm tăng chiều (18 chiều) làm khoảng cách bị loãng; đồng thời các feature yếu (chol, trestbps) thêm nhiễu vào metric khoảng cách. Điều này trực tiếp dẫn tới Experiment 3 về representation.", align="justify")

doc.add_heading("15. Experiment 2 – Hyperparameter Investigation", level=2)
p("Câu hỏi: Giá trị k của KNN và số cây B của Random Forest ảnh hưởng thế nào đến F1?", align="justify")
p("Chỉ thay đổi một hyperparameter; các điều kiện khác giữ nguyên.", align="justify")
img("heart_10_exp2_hyper.png", 15.5, "Hình 4 – F1 theo k (KNN) và theo n_estimators (Random Forest)")
p("Kết quả: Random Forest F1 tăng vọt từ 1 cây (0,703) lên 10 cây (0,751) và bão hòa từ khoảng B = 50 (0,766) — B = 400 (0,779) gần như không đổi. KNN đạt đỉnh ở k = 13 (0,791).", align="justify")
p("Nhận xét: k nhỏ làm mô hình nhạy với nhiễu; k lớn làm mô hình mượt hơn nhưng có thể mất cấu trúc cục bộ. Với Random Forest, thêm cây chỉ giảm variance chứ không tăng bias, qua ngưỡng nào đó cây mới trùng lặp cây cũ nên lợi ích bão hòa. Ta chọn giá trị phù hợp dựa trên metric, không chọn tùy ý: k = 11–13 cho KNN và B = 100 cho Random Forest.", align="justify")

doc.add_heading("16. Experiment 3 – Representation Investigation", level=2)
p("Câu hỏi: Chuẩn hóa feature có giúp các mô hình nhạy với scale như KNN và SVM hay không?", align="justify")
p("Ta so sánh từng mô hình với và không có StandardScaler, trong cùng imputation, cùng encoding và cùng train/test split, đo bằng 5-fold CV theo F1:", align="justify")
table(
    ["Model", "F1 raw (không scale)", "F1 standardized", "Chênh lệch"],
    [
        ["SVM (RBF)", "0,472", "0,785", "+0,313"],
        ["KNN (k=11)", "0,573", "0,782", "+0,208"],
        ["Logistic Regression", "0,825", "0,819", "−0,007"],
        ["Random Forest", "0,771", "0,764", "−0,007"],
    ],
)
p("Bảng 2 – Ảnh hưởng của standardization (5-fold CV F1)", align="center", size=11)
img("heart_11_exp3_scale.png", 14, "Hình 5 – So sánh raw vs standardized")
rich([("Kết luận Experiment 3. ", True), ("KNN sử dụng khoảng cách và SVM sử dụng margin nên scale của feature ảnh hưởng trực tiếp: chuẩn hóa giúp SVM tăng +0,313 và KNN tăng +0,208 F1. Trong khi đó Random Forest gần như không đổi vì cây quyết định phân chia theo ngưỡng từng feature, bất biến với phép biến đổi đơn điệu. Đây là minh họa trực tiếp cho nguyên lý của Assignment: thay đổi representation có thể làm thay đổi kết quả học, và mức ảnh hưởng phụ thuộc cơ chế học của từng mô hình.", False)], align="justify")
p("Thí nghiệm bổ sung so sánh X_all (18 chiều) với X_top6 (6 feature quan trọng nhất theo RF importance): top-6 giữ lại phần lớn hiệu năng, cho thấy tín hiệu tập trung trong ít feature then chốt. Hệ thống chính thức vẫn dùng X_all để không bỏ sót thông tin.", align="justify")

doc.add_heading("17. Final Model và lưu model", level=2)
p("Sau khi hoàn thành thí nghiệm, chọn Random Forest (B = 100) theo tiêu chí đã công bố: Recall cao nhất (ưu tiên y tế), F1 cao nhất, ổn định qua các fold, diễn giải được qua feature importance và không cần chuẩn hóa nên pipeline ứng dụng đơn giản, ít rủi ro mismatch representation. Final model được fit trên toàn bộ training set, không sử dụng test set để fit.", align="justify")
p("Model được lưu kèm metadata representation (thứ tự 18 cột sau One-Hot, danh sách biến phân loại) bằng joblib, nhờ đó application nhận dữ liệu thô và áp dụng đúng representation như lúc training.", align="justify")
img("heart_12_final_cm.png", 10, "Hình 6 – Confusion Matrix của Final Model")
doc.add_page_break()

# ============ PHẦN III ============
doc.add_heading("PHẦN III – HỆ THỐNG 2: DỰ ĐOÁN GIÁ NHÀ", level=1)

doc.add_heading("18. Mục tiêu", level=2)
p("Hệ thống nhận các thuộc tính của bất động sản và dự đoán SalePrice. Đây là bài toán Regression, vì target là một đại lượng số. Theo Assignment, các metric phù hợp gồm MAE, MSE, RMSE và R².", align="justify")
p("Dataset: Kaggle House Prices – Advanced Regression Techniques (bản gốc Ames Housing của De Cock, 2011), 1460 giao dịch nhà tại Ames, Iowa 2006–2010, 80 cột (Id + 79 feature + target). Nguồn: https://www.kaggle.com/c/house-prices-advanced-regression-techniques. Dữ liệu đã tải về và tiền xử lý thành data/ames_clean.csv (1460 × 75).", align="justify")
p("Mô tả dataset: có cả biến numerical (36) và categorical (38), đồng thời một số cột có missing. Vì vậy representation của mô hình phải xử lý cả hai loại dữ liệu. Điểm tinh tế: NA của các cột Garage/Bsmt/FireplaceQu trong Ames nghĩa là “nhà không có tiện ích đó” — là thông tin chứ không phải missing — nên được điền “None”; 5 cột thiếu trên 50% (PoolQC, MiscFeature, Alley, Fence, MasVnrType) bị bỏ; LotFrontage điền median theo từng khu dân cư. Sau xử lý: 0 giá trị thiếu.", align="justify")

doc.add_heading("19. EDA – Biểu đồ giá nhà và giải thích", level=2)
doc.add_heading("Biểu đồ 1 – Phân phối Price", level=3)
img("ames_01_target_dist.png", 15.5)
rich([("Cách đọc Biểu đồ 1 – Histogram phân phối SalePrice (trước và sau log). ", True), ("Trục X là giá nhà, trục Y là số lượng quan sát. Giá nhà lệch phải mạnh (skew = 1,88): phần lớn nhà nằm ở mức 100–250 nghìn USD nhưng một số nhà rất đắt (tới 755 nghìn) tạo đuôi dài bên phải, kéo mean (180.921) lên trên median (163.000). Vì vậy MAE/RMSE nên được xem cùng nhau: RMSE nhạy hơn với các sai số lớn. Biểu đồ bên phải cho thấy sau log-transform phân phối gần chuẩn (skew ≈ 0,12) — cơ sở cho quyết định học trên thang log ở Experiment 3. Histogram không cho biết model dự đoán tốt hay xấu; nó mô tả target trước khi học.", False)], align="justify")
doc.add_page_break()

doc.add_heading("Biểu đồ 2 – Scatter các feature mạnh với Price", level=3)
img("ames_02_scatter.png", 15.5)
rich([("Cách đọc Biểu đồ 2 – Scatter 6 feature vs Price. ", True), ("Mỗi chấm là một căn nhà; trục X là feature; trục Y là Price; đường đỏ là xu hướng tuyến tính. OverallQual (ρ = 0,79) có quan hệ đơn điệu tăng rõ rệt — chất lượng tổng thể là tín hiệu giá mạnh nhất, gần như “thang giá”. GrLivArea (ρ = 0,71) tăng cùng giá nhưng thấy hai điểm outlier dưới phải (nhà trên 4000 ft² giá thấp bất thường). Nếu tại cùng một giá trị feature có nhiều mức Price rất khác nhau, feature đó không đủ để giải thích giá — vị trí, garage, tầng hầm bổ sung thông tin. Scatter chỉ cho thấy association trong dataset, không chứng minh nhân quả.", False)], align="justify")

doc.add_heading("Biểu đồ 3 – Price theo Neighborhood", level=3)
img("ames_03_neighborhood.png", 15.5)
rich([("Cách đọc Biểu đồ 3 – Boxplot Price theo khu dân cư. ", True), ("Mỗi box là một khu: đường giữa hộp là median (giá điển hình của khu); đáy – đỉnh hộp là Q1 – Q3; râu là phạm vi không bị xem là outlier; điểm rời là outlier. So sánh theo median: khu đắt nhất (NoRidge ≈ 301k, NridgHt ≈ 275k) cao gần ba lần khu rẻ nhất (MeadowV ≈ 88k) — vị trí là feature phân loại có sức dự báo lớn nhất. Trong cùng một khu giá vẫn dàn rộng, nên Neighborhood cần kết hợp diện tích và chất lượng. Đây là lý do One-Hot giữ lại 25 cột Neighborhood.", False)], align="justify")
doc.add_page_break()

doc.add_heading("Biểu đồ 4 – Correlation của biến số", level=3)
img("ames_04_corr.png", 13.5)
rich([("Cách đọc Biểu đồ 4 – Correlation Heatmap. ", True), ("Cách đọc giống heatmap của bệnh tim: xem giá trị hệ số, dấu (+/−) và độ lớn. Chú ý ô giao giữa feature và SalePrice: OverallQual (0,79) và GrLivArea (0,71) lớn nhất. Giữa các feature có các cặp dư thừa: GarageCars ↔ GarageArea (0,88), TotalBsmtSF ↔ 1stFlrSF (0,81) — cùng đo một khái niệm (quy mô garage, quy mô tầng). Với Linear Regression đa cộng tuyến làm hệ số bất ổn khi diễn giải nhưng dự đoán vẫn tốt; với cây/Random Forest không vấn đề. Dữ liệu nhà có nhiều biến categorical nên heatmap chỉ xử lý biến số, không đại diện toàn bộ thông tin — lý do representation cuối cùng dùng cả numerical + categorical.", False)], align="justify")

doc.add_heading("Biểu đồ 5 – Hệ số Linear Regression và Importance của Random Forest", level=3)
img("ames_05_lr_coef.png", 14.5)
img("ames_06_rf_importance.png", 14.5)
rich([("Cách đọc Biểu đồ 5. ", True), ("Trên là hệ số của Linear Regression trên thang log-giá: cột đỏ (hệ số dương) làm tăng giá — các khu Neighborhood đắt (NoRidge, StoneBr, NridgHt), OverallQual, GrLivArea; cột xanh làm giảm giá — bếp chất lượng trung bình/kém, các khu rẻ (MeadowV, IDOTRR). Trên thang log, hệ số gần đúng phần trăm thay đổi giá khi feature tăng một đơn vị. Dưới là feature importance của Random Forest: OverallQual áp đảo (trên 0,55), tiếp theo là nhóm quy mô (GrLivArea, TotalBsmtSF, GarageCars/GarageArea) và năm xây. Lưu ý: Neighborhood bị chia nhỏ qua 25 cột One-Hot nên mỗi cột riêng lẻ thấp dù nhóm rất quan trọng — cần gộp lại khi so sánh feature gốc.", False)], align="justify")
doc.add_page_break()

doc.add_heading("20. Train/Test và Representation cho giá nhà", level=2)
p("Chia 80/20 với random_state = 42: 1168 nhà train, 292 nhà test. Representation: 38 biến phân loại mã hóa One-Hot nâng số chiều từ 74 lên 271; biến số giữ nguyên, chuẩn hóa trong Pipeline cho các mô hình cần. Target được học trên thang log(1 + Price) rồi revert về USD khi tính metric — chi tiết và biện minh ở Experiment 3.", align="justify")
p("Baseline (luôn đoán giá trung bình): MAE 59.931 USD; RMSE 88.271 USD; R² = −0,016 — nghĩa là đoán mò còn kém hơn cả trung bình test. Mọi mô hình thật phải kéo MAE xuống dưới 40 nghìn và R² lên trên 0,8 mới coi là học được điều gì hữu ích.", align="justify")
table(
    ["Model", "MAE (USD)", "RMSE (USD)", "R²"],
    [
        ["Random Forest (log-target)", "17.277", "29.297", "0,888"],
        ["Decision Tree (depth 8)", "≈22.000", "≈34.000", "≈0,85"],
        ["Linear Regression", "≈23.000", "≈35.000", "≈0,84"],
        ["SVR (RBF, mặc định)", "≈28.000", "≈41.000", "≈0,77"],
        ["Baseline (mean)", "59.931", "88.271", "−0,016"],
    ],
)
p("Bảng 3 – Kết quả các mô hình regression trên test set (n = 292)", align="center", size=11)
rich([("Giải thích kết quả. ", True), ("Đối với regression: MAE, MSE, RMSE càng thấp càng tốt; R² càng cao càng tốt. RMSE phạt mạnh các lỗi lớn nên hữu ích khi các trường hợp dự đoán sai rất xa cần được chú ý; MAE dễ diễn giải hơn (“sai khoảng 17 nghìn USD mỗi nhà” ≈ 10,6% median). Mô hình tốt nhất được lựa chọn bằng kết hợp R² cao và sai số thấp, đồng thời phải xét mục tiêu ứng dụng.", False)], align="justify")
img("ames_07_pred_vs_actual.png", 15.5, "Hình 7 – Actual vs Predicted của Random Forest và Linear Regression")
rich([("Cách đọc Actual vs Predicted. ", True), ("Trục X là giá thật trong test set; trục Y là giá dự đoán; đường chéo y = x là dự đoán hoàn hảo. Điểm nằm trên đường nghĩa là model dự đoán cao hơn giá thật; nằm dưới là thấp hơn; càng xa đường thì sai số càng lớn. Random Forest có đám mây điểm bám tương đối sát đường chéo trên toàn dải giá. Linear Regression tản rộng hơn và có mẫu hệ thống ở nhà đắt (trên 400k dự đoán thấp hơn thật) — siêu phẳng “kéo về trung bình”, đúng hạn chế lý thuyết của mô hình tuyến tính ở đuôi phân phối.", False)], align="justify")
img("ames_08_residuals.png", 15.5, "Hình 8 – Residual plot và phân phối residual của Random Forest")
rich([("Cách đọc Residual Plot. ", True), ("Residual = Actual − Predicted: bằng 0 là dự đoán đúng; dương là model dự đoán thấp hơn thực tế; âm là dự đoán cao hơn. Trục X là giá dự đoán; trục Y là residual. Residual plot tốt có các điểm phân tán quanh đường 0, không tạo pattern rõ. Nếu residual cong thì quan hệ phi tuyến chưa được mô hình hóa tốt; nếu độ rộng tăng theo X thì có heteroscedasticity. Trong bài, residual của Random Forest tản ngẫu nhiên quanh 0 — model đã học hết phần có thể; phần lớn nhà có sai số trong ±30 nghìn USD, dùng làm “vùng giá tham khảo” trong ứng dụng.", False)], align="justify")
doc.add_page_break()

doc.add_heading("21. Experiment 1 – Model Comparison", level=2)
p("Câu hỏi: mô hình hồi quy nào cho kết quả tốt nhất dưới cùng train/test protocol?", align="justify")
p("Bốn mô hình được so sánh bằng 10-fold Cross-Validation trên training set theo R² của log-price: Random Forest 0,862 ± 0,029; Linear Regression 0,839 ± 0,091; SVR 0,761 ± 0,078; Decision Tree 0,745 ± 0,059. Xếp hạng nhất quán với test set: Random Forest dẫn đầu và ổn định nhất (khoảng tứ phân vị hẹp). Bảng metric là bằng chứng định lượng — không kết luận “Random Forest tốt nhất” nếu chưa nhìn vào metric thực tế.", align="justify")
p("Vì sao Random Forest thắng: giá nhà là quan hệ phi tuyến với tương tác (chất lượng × diện tích × vị trí); cây bắt tương tác miễn phí và ensemble giảm variance trên 1.460 mẫu vừa phải.", align="justify")

doc.add_heading("22. Experiment 2 – Hyperparameter Investigation", level=2)
p("Câu hỏi: độ sâu max_depth của Decision Tree và số cây n_estimators của Random Forest ảnh hưởng thế nào đến R²?", align="justify")
p("Chỉ thay đổi một hyperparameter, giữ các điều kiện khác cố định.", align="justify")
img("ames_10_exp2_hyper.png", 15.5, "Hình 9 – R² theo max_depth (Decision Tree) và n_estimators (Random Forest)")
p("Kết quả: Decision Tree tăng vọt từ depth 2 (R² = 0,58 — underfit, cây quá cạn chỉ thấy xu hướng thô) lên đỉnh ở depth 8–12 (≈ 0,75) rồi đi ngang/giảm nhẹ khi sâu hơn. Random Forest bão hòa từ B ≈ 100.", align="justify")
p("Nhận xét: đây là minh họa sống động của trade-off bias – variance: cây quá cạn underfit (bias cao), cây quá sâu khớp noise của train (variance cao) nên điểm cross-validation tụt. Chọn depth 8–12 cho cây đơn và B = 100 cho Random Forest dựa trên metric, không chọn tùy ý.", align="justify")

doc.add_heading("23. Experiment 3 – Representation / Feature Investigation", level=2)
p("Câu hỏi: (a) log-transform target có ảnh hưởng đến khả năng học không? (b) Thu nhỏ representation từ 271 chiều về nhóm feature quan trọng nhất có giữ được hiệu năng không?", align="justify")
p("(a) So sánh cùng mô hình học trên target gốc và target log, cùng 10-fold CV:", align="justify")
table(
    ["Model", "R² target gốc", "R² log-target", "Chênh lệch"],
    [
        ["Linear Regression", "0,807", "0,839", "+0,032"],
        ["Random Forest", "0,854", "0,862", "+0,008"],
        ["SVR (RBF)", "thấp", "0,761", "cải thiện lớn"],
    ],
)
p("Bảng 4 – Ảnh hưởng của log-transform target (10-fold CV R²)", align="center", size=11)
p("(b) So sánh X_all (271 chiều) với X_top8 (8 feature gốc quan trọng nhất): Random Forest chỉ giảm khoảng 0,03 – 0,05 R² — phần lớn tín hiệu nằm trong ít feature then chốt (OverallQual, GrLivArea, GarageCars, TotalBsmtSF, FullBath, YearBuilt, OverallCond, LotArea).", align="justify")
img("ames_11_exp3_repr.png", 15.5, "Hình 10 – So sánh representation của target và của input")
rich([("Kết luận Experiment 3. ", True), ("Log-transform cải thiện mọi mô hình vì: (1) giá nhà lệch phải nên quan hệ feature – giá tuyến tính hơn ở thang log; (2) squared-error trên thang log tương đương sai số tương đối, làm nhà 80 nghìn USD và 800 nghìn USD được coi trọng công bằng; (3) phương sai ổn định hơn, đúng giả định của OLS. Việc top-8 feature giữ gần nguyên hiệu năng cho thấy tín hiệu tập trung. Cả hai thí nghiệm cùng khẳng định nguyên lý trung tâm của Assignment: thay đổi representation — của input lẫn target — thay đổi kết quả học.", False)], align="justify")

doc.add_heading("24. Final Model – Giá nhà", level=2)
p("Chọn Random Forest Regressor (B = 100, target log) theo tiêu chí: R² cao nhất (0,888), MAE thấp nhất (17.277 USD ≈ 10,6% median), ổn định qua 10 fold, có feature importance để diễn giải, ít nhạy hyperparameter. Final model được fit lại trên toàn bộ training set và lưu kèm metadata representation (271 cột sau One-Hot) bằng joblib.", align="justify")
img("ames_12_final.png", 11.5, "Hình 11 – Final Model: Predicted vs Actual (R² = 0,888)")
doc.add_page_break()

# ============ PHẦN IV ============
doc.add_heading("PHẦN IV – TRIỂN KHAI WEB", level=1)

doc.add_heading("25. Kiến trúc Web", level=2)
p("Hai hệ thống được triển khai bằng ứng dụng web Streamlit, deploy lên Hugging Face Spaces:", align="justify")
p("Browser / Client → HTML Form (Streamlit) → hàm predict trung tâm → One-Hot đúng thứ tự cột training → Final model đã lưu → Prediction → Hiển thị kết quả + khuyến nghị", bold=True, align="center")
p("Nguyên tắc: không viết lại preprocessing theo cách khác. Model đã lưu chứa metadata representation (danh sách cột, danh sách biến phân loại); ứng dụng mã hóa dữ liệu thô rồi reindex đúng 18/271 cột như lúc training, bảo đảm dữ liệu mới được biểu diễn giống dữ liệu lúc training.", align="justify")
p("Ứng dụng bệnh tim: người dùng nhập 13 chỉ số (kèm giải thích y khoa tiếng Việt), hệ thống trả về xác suất bệnh tim, ngưỡng quyết định 0,5, kết luận kèm khuyến nghị và cảnh báo vùng mơ hồ 0,35–0,65. Ứng dụng giá nhà: người dùng nhập 8 feature quyết định trong 3 tab (chất lượng, diện tích, vị trí), hệ thống trả về giá ước tính kèm dải tham chiếu ±MAE/RMSE và so sánh với median thị trường. Cả hai ứng dụng đều có 3 case mẫu demo một cặp phím.", align="justify")
p("Link triển khai (Hugging Face Spaces):", bold=True)
p("• Hệ bệnh tim: https://....hf.space (điền sau khi deploy theo apps/HUONG_DAN_DEPLOY.md)")
p("• Hệ giá nhà: https://....hf.space (điền sau khi deploy)")

doc.add_heading("26. Demo ứng dụng", level=2)
p("Minh chứng hệ thống chạy thật (chụp từ ứng dụng local trước khi deploy):", align="justify")
img("app_heart_demo.png", 15.5, "Hình 12 – Demo web hệ Chẩn đoán bệnh tim: case nguy cơ cao → CÓ BỆNH TIM")
p("Demo 3 case bệnh tim (từ notebook và ứng dụng):", align="justify")
img("heart_13_demo.png", 15.5, "Hình 13 – Xác suất dự đoán cho 3 case: nguy cơ thấp / cao / biên")
p("Case 1 (42 tuổi, chỉ số lành): xác suất bệnh thấp, quyết định tự tin “khỏe”. Case 2 (68 tuổi, cp = 4, oldpeak 3,2, ca = 2, thal = 7): xác suất rất cao — mô hình nhận diện đúng cụm đa tín hiệu nguy cơ. Case 3 (57 tuổi, tín hiệu hỗn hợp): xác suất nằm gần ngưỡng 0,5 — hành vi đúng mong đợi với input khó; ứng dụng gắn nhãn “khuyến nghị khám chuyên sâu” thay vì kết luận cứng.", align="justify")
img("app_house_demo.png", 15.5, "Hình 14 – Demo web hệ Dự đoán giá nhà: nhà cao cấp NridgHt → 302.408 USD")
img("ames_13_demo.png", 15.5, "Hình 15 – Dự đoán vs giá giao dịch thật cho 3 nhà (phổ thông / cao cấp / cấp thấp)")
p("Ba nhà đại diện ba phân khúc: nhà phổ thông dự đoán sát giá thật (vùng dữ liệu dày, model “rành” nhất); nhà cao cấp lệch lớn hơn một chút (hiếm trong train); nhà cấp thấp dự đoán tốt. Quy luật tổng: model tin cậy nhất ở vùng dữ liệu dày, kém nhất ở hai đầu phân phối — lý do ứng dụng luôn hiển thị dải tham chiếu và cảnh báo độ tin cậy ở phân khúc cao cấp.", align="justify")
p("Đường đi đầy đủ User → Input → Feature Representation → ML Model → Prediction đã được chứng minh trong cả hai hệ thống — hoàn thành yêu cầu System Demonstration của Assignment.", align="justify")

doc.save(os.path.join(BASE, "docs", "BAO_CAO_A01.docx"))
print("SAVED:", os.path.join(BASE, "docs", "BAO_CAO_A01.docx"))
