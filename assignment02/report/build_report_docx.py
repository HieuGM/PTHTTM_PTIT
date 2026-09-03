# -*- coding: utf-8 -*-
"""Build BAO_CAO_A02.docx — báo cáo Assignment 02 theo Report Template của đề.
Chạy: python build_report_docx.py  (từ assignment02/report/)
"""
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

FIG = Path(__file__).parent.parent / "figures"
OUT = Path(__file__).parent / "BAO_CAO_A02.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

def h1(t):
    doc.add_heading(t, level=1)

def h2(t):
    doc.add_heading(t, level=2)

def p(text, bold=False, italic=False, size=12, center=False):
    par = doc.add_paragraph()
    if center:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return par

def bullet(text):
    par = doc.add_paragraph(style="List Bullet")
    par.add_run(text).font.size = Pt(12)

def table(rows, header=True, widths=None, size=10):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_i, row in enumerate(rows):
        for c_i, cell in enumerate(row):
            cell_obj = t.cell(r_i, c_i)
            cell_obj.text = ""
            run = cell_obj.paragraphs[0].add_run(str(cell))
            run.font.size = Pt(size)
            if header and r_i == 0:
                run.bold = True
    doc.add_paragraph()
    return t

def fig(name, caption, width_cm=15.5):
    path = FIG / name
    if path.exists():
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(str(path), width=Cm(width_cm))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(10)
    else:
        p(f"[Hình: {name} — {caption}]", italic=True, size=10, center=True)

# ============================ TRANG BÌA ============================
p("", size=12); p("", size=12)
p("HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG", bold=True, size=16, center=True)
p("POSTS AND TELECOMMUNICATIONS INSTITUTE OF TECHNOLOGY", size=12, center=True)
p("", size=12)
p("INTELLIGENT SYSTEM DEVELOPMENT", bold=True, size=20, center=True)
p("ASSIGNMENT 02", bold=True, size=24, center=True)
p("From Data Representation to Deployable Intelligent Systems", italic=True, size=14, center=True)
p("", size=12); p("", size=12)
p("Sinh viên: ............................................", size=13)
p("MSSV: ............................................", size=13)
p("Lớp: ............................................", size=13)
p("Giảng viên: PGS. TS. Trần Đình Quế", size=13)
p("Học kỳ: I.2026", size=13)
doc.add_page_break()

# ============================ 1. EXECUTIVE SUMMARY ============================
h1("1. Executive Summary")
p("Bài tập xây dựng ba hệ thống thông minh trên ba bộ dữ liệu Kaggle khác nhau, cùng tuân "
  "theo một quy trình phát triển thống nhất: Data → Understand → Clean → Represent → Learn → "
  "Evaluate → Persist → Deploy. Mỗi hệ thống được triển khai đầy đủ REST API (FastAPI), web app "
  "(Streamlit) và mobile client (HTML mobile-first gọi API).")

table([
    ["Ứng dụng", "Nhiệm vụ", "Representation chính", "Model chọn", "Kết quả test", "Deployment"],
    ["Diabetes", "Binary classification",
     "Feature matrix 8 chiều (impute + scale)", "Random Forest (200)",
     "Acc 0.74 · AUC 0.81", "FastAPI + Streamlit + Mobile"],
    ["House Price", "Regression",
     "17 numeric (scale) + 70 one-hot zipcode, log-target", "Gradient Boosting",
     "R² 0.90 · MAE $68k", "FastAPI + Streamlit + Mobile"],
    ["Customer Behavior", "Multi-class (9 lớp)",
     "RFM 5D ⊕ TF-IDF 4000D (text embedding)", "Linear SVM (text-linear)",
     "macro-F1 0.37 · Acc 0.70", "FastAPI + Streamlit + Mobile"],
])

p("Điểm nhấn representation (trọng tâm Lecture 02):")
bullet("App 1: phát hiện 0 sinh lý là missing trá hình (Insulin thiếu 48.7%) — median impute học từ train.")
bullet("App 2: zipcode (70 vùng) one-hot + log(1+price) chuẩn hóa skew 4.02 → 0.43.")
bullet("App 3: pipeline đầy đủ Comment → Tokens → Token IDs → TF-IDF embedding; thí nghiệm đối "
       "chứng cho thấy thêm text nâng macro-F1 của Random Forest từ 0.163 lên 0.291.")

# ============================ 2. CONNECTION LECTURE 02 ============================
h1("2. Connection to Lecture 02 — Data Representation")
p("Lecture 02 đặt nguyên lý: Real-world data → numerical representation → computational model. "
  "Ba ứng dụng minh họa ba dạng biểu diễn khác nhau của cùng một nguyên lý:")
table([
    ["Ứng dụng", "Raw Data", "ML Representation", "Model input"],
    ["Diabetes", "CSV — 8 chỉ số lâm sàng/bệnh nhân", "Feature matrix X ∈ R^(768×8) sau impute + standardize", "B × 8 float64"],
    ["House Price", "CSV — 19 đặc điểm/nhà (có zipcode)", "Encoded matrix 17 scaled + 70 one-hot; y = log(1+price)", "B × 87 float64"],
    ["Customer", "CSV giao dịch + mô tả sản phẩm (text)", "RFM scaled ⊕ TF-IDF sparse; E ∈ R^(B×4000)", "sparse B × 4005; demo B×T×384"],
])
p("Trả lời 9 câu hỏi representation chung:")
table([
    ["#", "Câu hỏi", "Diabetes", "House", "Customer"],
    ["1", "1 hàng là gì?", "1 bệnh nhân", "1 giao dịch bán nhà", "1 dòng giao dịch → gộp thành 1 khách"],
    ["2", "1 cột là gì?", "1 chỉ số lâm sàng", "1 đặc điểm nhà", "thuộc tính đơn / sản phẩm / khách"],
    ["3", "Input features?", "8 cột", "17 cột (sau FE)", "5 RFM + basket text"],
    ["4", "Target?", "Outcome (0/1)", "price (USD)", "interest (9 nhóm)"],
    ["5", "Numerical?", "8/8", "16/17", "5 RFM"],
    ["6", "Categorical?", "0", "zipcode (70)", "interest (target)"],
    ["7", "Encode categorical?", "—", "one-hot zipcode", "label interest; text → TF-IDF"],
    ["8", "Feature dimension cuối?", "8", "87", "4005"],
    ["9", "Model input shape?", "B × 8", "B × 87", "B × 4005 sparse"],
])

# ============================ 3-5. APP 1 ============================
doc.add_page_break()
h1("3. Application 1 — Diabetes Prediction")

h2("3.1 Problem Description")
p("Mục tiêu: dự đoán một bệnh nhân có mắc tiểu đường type 2 hay không từ 8 chỉ số lâm sàng "
  "(số lần mang thai, glucose, huyết áp, độ dày da, insulin, BMI, hệ số di truyền, tuổi). "
  "X = patient features ∈ R^8, y = diabetes class ∈ {0,1}. Ứng dụng hỗ trợ sàng lọc sớm, "
  "ưu tiên xét nghiệm HbA1c xác nhận. Dataset: Pima Indians Diabetes (Kaggle, "
  "https://www.kaggle.com/datasets/uciml/pima-indians-diabetes-database), 768 quan sát × 9 cột.")

h2("3.2 Data Understanding & Quality")
bullet("Không có NaN chính thức, NHƯNG 5 cột sinh lý có giá trị 0 vô lý (không ai sống với glucose = 0): "
       "Insulin 374 (48.7%), SkinThickness 227 (29.6%), BloodPressure 35, BMI 11, Glucose 5 → 0 chính là missing trá hình.")
bullet("Không trùng lặp; không invalid ngoài zero; outlier (Insulin 846...) là ca bệnh thật → GIỮ, dùng median impute kháng outlier.")
bullet("Mất cân bằng nhẹ 500/268 (65/35) → stratify split + ưu tiên Recall.")
fig("diabetes_01_missing.png", "Hình 3.1 — Missing values sau khi chuẩn hóa 0 sinh lý → NaN")

h2("3.3 Cleaning & Representation")
p("Quy trình: CSV → ép 0 vô lý → NaN → median impute (fit trên train) → StandardScaler → "
  "x ∈ R^8. N = 768, d = 8; model input batch B × 8 float64. Representation Contract đầy đủ "
  "(shape, dtype, range, encoding, missing, split, model input) được ghi trong notebook mục 12.")

h2("3.4 EDA (trích)")
fig("diabetes_04_feature_dist.png", "Hình 3.2 — Phân phối 8 feature theo lớp: Glucose tách lớp rõ nhất")
p("Observation: Glucose phân phối nhóm bệnh dịch phải mạnh. Interpretation: glucose là marker "
  "sinh lý trực tiếp của bệnh. ML implication: dự kiến Glucose chiếm feature importance cao nhất "
  "— kiểm chứng bằng RF importance (Glucose ≈ 0.33, cao nhất).", italic=True)
fig("diabetes_09_test_eval.png", "Hình 3.3 — Final model: Confusion Matrix + ROC 5 model trên test")

h2("3.5 Models & Evaluation (5 model)")
table([
    ["Model (VAL n=115)", "Accuracy", "Precision", "Recall", "F1"],
    ["KNN (k=13)", "0.765", "0.686", "0.600", "0.640"],
    ["Random Forest", "0.757", "0.667", "0.600", "0.632"],
    ["SVM (RBF)", "0.748", "0.657", "0.575", "0.613"],
    ["Logistic Regression", "0.713", "0.595", "0.550", "0.571"],
    ["Decision Tree", "0.696", "0.576", "0.475", "0.521"],
    ["Baseline (majority)", "0.652", "0.000", "0.000", "0.000"],
])
p("Test (n=116, Random Forest — model deploy): Accuracy 0.741, Precision 0.690, Recall 0.488, "
  "F1 0.571, ROC-AUC 0.812. Metric quan trọng nhất cho sàng lọc y tế là RECALL (bỏ sót người "
  "bệnh = FN là lỗi nặng nhất); AUC 0.81 cho thấy khả năng xếp hạng tốt bất kể ngưỡng. KNN nhỉnh "
  "hơn RF trên VAL-F1 0.008 nhưng RF được chọn để triển khai vì diễn giải được (feature "
  "importance), inference nhanh độc lập train set, ổn định theo tham số.")

h2("3.6 Deployment")
p("Pipeline (imputer + scaler + RF + metadata) lưu 1 file diabetes_pipeline.joblib. FastAPI "
  "POST /predict nhận 8 chỉ số, áp đúng preprocessing đã lưu, trả {prediction, confidence}. "
  "Web Streamlit + mobile HTML gọi cùng API đó — cùng preprocessing cho training và serving.")
fig("web_diabetes_demo.png", "Hình 3.4 — Web app: nhập chỉ số → dự đoán qua API")
fig("mobile_diabetes_02_result.png", "Hình 3.5 — Mobile app: dự đoán CÓ TIỂU ĐƯỜNG 90.0% (gọi API thật)", width_cm=8)

# ============================ 4-6. APP 2 ============================
doc.add_page_break()
h1("4. Application 2 — House Price Prediction")

h2("4.1 Problem Description")
p("Dự đoán giá bán nhà tại Quận King (Seattle, WA) từ đặc điểm nhà. X = house features, "
  "y = price (USD) — regression, khác bản chất App 1: target liên tục, metric là khoảng cách "
  "(MAE/RMSE) và % biến thiên giải thích (R²). Dataset: House Sales in King County (Kaggle, "
  "https://www.kaggle.com/datasets/harlfoxem/housesalesprediction), 21,613 giao dịch × 21 cột "
  "(05/2014–05/2015).")

h2("4.2 Data Understanding & Cleaning")
bullet("id trùng 177 (nhà bán lại) → giữ giao dịch gần nhất: giá mới nhất + tránh cùng nhà rơi vào train và test.")
bullet("bedrooms = 33 với 1,620 sqft là lỗi nhập → sửa về median nhóm diện tích tương đương.")
bullet("Không missing; outlier giá cao (đến $7.7M) là THẬT → giữ, xử lý bằng log-target.")
fig("house_01_price_dist.png", "Hình 4.1 — Price lệch phải (skew 4.02) → log(1+price) gần chuẩn (0.43)")

h2("4.3 Representation")
p("17 biến numeric (trong đó 2 feature engineered: house_age = 2015 − yr_built, renovated ∈ {0,1}) "
  "được StandardScaler; zipcode — categorical 70 vùng — one-hot thành 70 cột nhị phân. "
  "X ∈ R^(21536×87), y = log(1+price). Ví dụ categorical→numerical: zipcode 98103 → "
  "[0,…,1,…,0] (70 chiều) — model học weight riêng từng vùng mà không giả định thứ tự số "
  "của mã vùng. Batch: B × 87 float64.")

h2("4.4 EDA (trích)")
fig("house_05_zipcode.png", "Hình 4.2 — Giá trung vị theo zipcode: vùng đắt nhất ~7× vùng rẻ nhất")
fig("house_07_pred_vs_actual.png", "Hình 4.3 — Dự đoán vs thực tế (VAL): sát đường chéo ở vùng dày dữ liệu")
p("Observation: điểm tản mạnh và lệch xuống trên $2M. Interpretation: nhà siêu sang hiếm trong "
  "train, log-target bình phương sai số tương đối. ML implication: web app hiển thị dải tham "
  "chiếu ±12% và cảnh báo độ tin cậy ở phân khúc cao cấp.", italic=True)

h2("4.5 Models & Evaluation (5 model)")
table([
    ["Model (VAL n=3230)", "MAE", "RMSE", "R²", "Fit+Infer (s)"],
    ["Gradient Boosting", "$70,624", "$144,357", "0.871", "0.05"],
    ["Random Forest", "$71,558", "$154,051", "0.854", "0.18"],
    ["Decision Tree (d=10)", "$89,405", "$165,006", "0.832", "0.02"],
    ["Ridge (α=1)", "$80,439", "$239,225", "0.647", "0.03"],
    ["Linear Regression", "$80,608", "$245,997", "0.627", "0.03"],
    ["Baseline (mean)", "$222,729", "$409,369", "−0.034", "0.01"],
])
p("Test (n=3231, Gradient Boosting): MAE $68,225, RMSE $116,374, R² 0.897. MAE là metric chính "
  "cho ứng dụng định giá — 'trung bình lệch $68k' diễn giải trực tiếp cho khách hàng; R² 0.90 "
  "nghĩa là model giải thích 90% biến thiên giá. Error analysis: sai số trung vị ~12% nhà phổ "
  "thông (<$1M) vs ~19% nhà cao cấp — model tin cậy nhất nơi dữ liệu dày.")

h2("4.6 Deployment")
p("ColumnTransformer (scaler + one-hot) + GB lưu 1 file house_pipeline.joblib. API POST /predict "
  "nhận 16 trường, tự tính house_age/renovated, dự đoán log → expm1 → USD, trả "
  "{predicted_price, confidence_range}.")
fig("web_house_demo.png", "Hình 4.4 — Web app: định giá nhà 1340 sqft zipcode 98103")
fig("mobile_house_02_result.png", "Hình 4.5 — Mobile: giá $478,771 + dải tham chiếu", width_cm=8)

# ============================ 5-7. APP 3 ============================
doc.add_page_break()
h1("5. Application 3 — E-commerce Customer Behavior & Interest")

h2("5.1 Problem Description")
p("Phân tích hành vi và phát hiện sở thích khách hàng từ dữ liệu giao dịch e-commerce. Bài toán "
  "đặt thành supervised: X = RFM hành vi + văn bản giỏ hàng, y = interest — category chiếm doanh "
  "thu lớn nhất của khách (9 lớp sau khi gộp nhóm hiếm). Dataset: Online Retail (Kaggle, "
  "https://www.kaggle.com/datasets/carrie1/ecommerce-data) — 55,263 giao dịch UK 2010–2011, "
  "1,033 khách có ID sau làm sạch.")

h2("5.2 Data Cleaning (từ transaction về customer)")
bullet("Bỏ 834 dòng lỗi parse; bỏ 20,889 dòng không CustomerID (bài toán cấp khách hàng); bỏ 1,002 đơn hủy 'C…'; bỏ Quantity/UnitPrice ≤ 0; khử 5,268 dòng trùng (InvoiceNo, StockCode); Description thiếu → 'UNKNOWN'.")
bullet("55,263 → 33,417 giao dịch sạch; outlier chi tiêu ($280k max) là khách B2B thật → GIỮ làm tín hiệu phân khúc.")
bullet("Tạo 9 category sản phẩm bằng keyword rules trên Description (home_decor, kitchen_dining, …) — tái lập được.")

h2("5.3 Customer Representation — trung tâm của app")
p("Transactions → Customer profile: RFM (recency_days, frequency, monetary→log, total_items→log, "
  "avg_order_value) + basket_text (gộp toàn bộ Description khách đã mua thành 1 văn bản giỏ hàng).")
p("Pipeline text (yêu cầu bắt buộc của đề):", bold=True)
table([
    ["Bước", "Biến đổi", "Ví dụ"],
    ["1. Tokens", "tách từ + n-gram (1,2)", '"white hanging heart t-light holder" → 5 token'],
    ["2. Token IDs", "vocabulary V học TỪ TRAIN, mỗi token ↦ integer", "heart → 1042; holder → 2031"],
    ["3. TF-IDF weights", "tf × log(N/df) — từ hiếm giá trị hơn", "w(heart,d) = 0.031"],
    ["4. Embedding", "mỗi khách ↦ vector thưa R^|V|", "E ∈ R^(B×4000) sparse"],
])
p("B, T, d: B = số khách trong batch (train 644 / val 153 / test 155); T = chiều dài văn bản giỏ "
  "hàng (chỉ tồn tại ở dạng sequence); d = 4000 (TF-IDF, 1 chiều/token). TF-IDF là bag-of-words "
  "nên không giữ T — notebook mục 18.6 đối chiếu thêm neural embedding MiniLM: mỗi token ↦ "
  "R^384, sequence T×384, batch B×T×384 (đúng dạng E ∈ R^(B×T×d) của Slide 02), mean-pooling về "
  "B×384.")
bullet("Chống leakage: vocabulary + idf chỉ fit trên train; category revenue share KHÔNG dùng làm feature (chúng sinh ra target bằng idxmax — đưa vào X sẽ được accuracy 100% giả).")

h2("5.4 EDA (trích)")
fig("ecom_02_categories.png", "Hình 5.1 — Giao dịch & doanh thu theo category: home_decor, kitchen_dining dẫn đầu")
fig("ecom_03_rfm.png", "Hình 5.2 — Phân phối RFM: đông khách mua 1 lần; monetary lệch phải mạnh")

h2("5.5 Models & Evaluation (6 model)")
table([
    ["Model (VAL n=153, macro-F1)", "Accuracy", "Precision(m)", "Recall(m)", "F1(macro)"],
    ["Linear SVM (text-linear)", "0.710", "0.401", "0.329", "0.325"],
    ["SVM (RBF)", "0.671", "0.386", "0.301", "0.309"],
    ["Random Forest", "0.652", "0.388", "0.277", "0.291"],
    ["Logistic Regression", "0.690", "0.427", "0.256", "0.276"],
    ["Decision Tree", "0.561", "0.219", "0.216", "0.200"],
    ["KNN (k=15)", "0.510", "0.243", "0.163", "0.150"],
    ["Baseline (majority)", "0.503", "0.000", "0.000", "0.058"],
])
p("Test (n=155, Linear SVM): accuracy 0.697, macro-F1 0.373. Macro-F1 là metric chính vì 9 lớp "
  "mất cân bằng — accuracy bị kéo bởi lớp đông (home_decor). Confusion matrix cho thấy nhầm lẫn "
  "tập trung ở các nhóm gần ngữ nghĩa (decor↔kitchen↔vintage_craft đều là đồ nhà cửa) — lỗi "
  "'hợp lý' chứ không ngẫu nhiên.")

h2("5.6 Tabular-only vs Tabular + Text (yêu cầu đề)")
table([
    ["Representation (macro-F1 VAL)", "Logistic Regression", "Random Forest", "SVM (RBF)"],
    ["Tabular-only (5 chiều RFM)", "0.083", "0.163", "0.109"],
    ["Tabular + Text (4005 chiều)", "0.276", "0.291", "0.309"],
])
p("Text nâng macro-F1 +0.11 đến +0.19 ở mọi model: cùng RFM, khách mua 'candle holder heart' và "
  "'teacup saucer' là hai interest khác nhau — thông tin chỉ có trong văn bản. Đây là bằng chứng "
  "thực nghiệm trực tiếp cho luận điểm của Lecture 02: representation quyết định thông tin "
  "available cho learning.")
fig("ecom_07_repr_cmp.png", "Hình 5.3 — Text cải thiện phân loại interest ở cả 3 model")

h2("5.7 Business Interpretation")
p("Hệ thống phân loại khách theo 9 nhóm sở thích từ hành vi + giỏ hàng. Ứng dụng: gợi ý sản "
  "phẩm theo nhóm, targeted promotion (giảm giá đúng nhóm hàng khách quan tâm), chăm sóc khách "
  "VIP (RFM Monetary cao), xếp chiến dịch marketing theo phân khúc. API trả kèm top-3 interest "
  "kèm xác suất để campaign chọn ngưỡng phù hợp.")

h2("5.8 Deployment")
p("scaler + TF-IDF (vocabulary/idf đã lưu) + Linear SVM + labels trong 1 joblib. API POST "
  "/predict nhận RFM + basket_text → transform (không fit lại) → interest + confidence + top-3.")
fig("web_customer_demo.png", "Hình 5.4 — Web app: dự đoán home_decor 96.2% + top-3 xác suất")
fig("mobile_customer_02_result.png", "Hình 5.5 — Mobile: interest home_decor, confidence 96.2%", width_cm=8)

# ============================ 6. CROSS-APP COMPARISON ============================
doc.add_page_break()
h1("6. Comparison of the Three Intelligent Systems")
table([
    ["Aspect", "Diabetes", "House Price", "Customer Behavior"],
    ["Problem type", "Binary classification", "Regression", "Multi-class classification (9)"],
    ["Observation", "1 bệnh nhân (phụ nữ Pima)", "1 giao dịch bán nhà", "1 khách (gộp từ giao dịch)"],
    ["Target", "Outcome {0,1}", "price (USD, log)", "interest (9 nhóm)"],
    ["Input representation", "8D imputed + scaled", "17D scaled + 70D one-hot", "5D RFM ⊕ 4000D TF-IDF"],
    ["Model input", "B × 8 float64", "B × 87 float64", "sparse B × 4005"],
    ["Data-quality issues", "0-trá hình (Insulin 48.7%)", "id trùng, bedrooms=33, skew giá", "thiếu CustomerID 38%, đơn hủy, trùng dòng"],
    ["Best model", "Random Forest", "Gradient Boosting", "Linear SVM (text)"],
    ["Main metric", "Recall / ROC-AUC", "MAE / R²", "Macro-F1"],
    ["Web deployment", "Streamlit + FastAPI", "Streamlit + FastAPI", "Streamlit + FastAPI"],
    ["Mobile deployment", "HTML mobile-first → API", "HTML mobile-first → API", "HTML mobile-first → API"],
    ["Main limitation", "nhỏ (768), Recall test 0.49", "kém tin ở nhà >$2M", "category rules thủ công, lớp hiếm yếu"],
])
p("Thảo luận so sánh:")
bullet("Dataset khác nhau thế nào? Tabular thuần (App 1), tabular nhiều biến + categorical vị trí (App 2), "
       "tabular + text phi cấu trúc (App 3) — ba mức phức tạp representation tăng dần.")
bullet("Preprocessing chung: median/log cho skew, scale cho model khoảng cách, chống leakage bằng fit-on-train. "
       "Riêng từng app: zero→NaN (y tế), one-hot zipcode + log-target (giá nhà), tokenization + TF-IDF (text).")
bullet("Target representation khác nhau vì bản chất quyết định khác nhau: nhãn rời rạc (chẩn đoán) cần Recall; "
       "giá liên tục cần sai số tiền (MAE); 9 lớp mất cân bằng cần macro-F1.")
bullet("Dễ deploy nhất: App 1 (input ít, model nhẹ). Nặng nhất: App 3 (TF-IDF 4000D + text processing, "
       "model SVM linear trên sparse matrix).")
bullet("Best model ≠ best-on-one-metric: App 1 chọn RF dù KNN nhỉnh hơn 0.008 VAL-F1 vì lý do triển khai "
       "(diễn giải, tốc độ, ổn định).")

# ============================ 7. DEPLOYMENT ARCHITECTURE ============================
h1("7. Deployment Architecture")
p("Cả ba hệ thống dùng cùng một kiến trúc inference:")
p("User Input → Validation (Pydantic) → Same Preprocessing (từ joblib) → Saved Model → "
  "JSON Response → Web/Mobile UI", center=True, italic=True)
bullet("API: FastAPI, endpoint POST /predict, validation Pydantic (giới hạn min/max từng trường), "
       "CORS mở cho web/mobile client, /health health-check, /docs Swagger UI.")
bullet("Web: Streamlit client KHÔNG chứa model — gọi API qua HTTP, cùng một nguồn sự thật.")
bullet("Mobile: HTML mobile-first (dark mode, viewport 390px) host ngay trong API Space tại /mobile "
       "— mobile là client của REST API đúng kiến trúc: Mobile UI → REST API → Preprocessing → "
       "Saved Model → Prediction → Mobile UI.")
bullet("Training ≠ Inference: model train trong notebook, lưu pipeline; deployed service chỉ load + "
       "transform + predict. Không fit scaler/encoder/imputer trên dữ liệu user (đề cảnh báo data leakage).")
fig("api_swagger_diabetes.png", "Hình 7.1 — Swagger UI API Diabetes: POST /predict, schema Pydantic hiện đầy đủ")

# ============================ 8. DISCUSSION QUESTIONS ============================
doc.add_page_break()
h1("8. Discussion Questions (15 câu chung + 6 câu e-commerce)")
h2("8.1 Mười lăm câu chung (trả lời gọn cho cả 3 app)")
qa = [
    ("1. Một observation là gì?", "A1: 1 bệnh nhân. A2: 1 giao dịch bán nhà. A3: 1 khách hàng (gộp từ nhiều giao dịch)."),
    ("2. Raw representation?", "A1: dòng CSV 9 cột. A2: dòng CSV 21 cột. A3: dòng giao dịch 8 cột kèm Description text."),
    ("3. Final numerical representation?", "A1: X∈R^(768×8) scaled. A2: X∈R^(21536×87). A3: sparse B×4005 = RFM⊕TF-IDF."),
    ("4. Chiều của feature matrix nghĩa gì?", "Mỗi chiều = 1 feature số học sau transform (A2 gồm 70 chiều one-hot zipcode; A3 gồm 4000 chiều từ vựng)."),
    ("5. Feature nào cần encode?", "A1: không (toàn numeric). A2: zipcode → one-hot. A3: text → token IDs → TF-IDF."),
    ("6. Feature nào cần normalization?", "A1: cả 8 (scale). A2: 17 numeric. A3: 5 RFM (log trước khi scale)."),
    ("7. Mất thông tin gì khi represent?", "A1: cơ chế missing (0 = không đo vs thật = 0). A2: thứ bậc vùng lân cận (one-hot coi mọi zipcode độc lập). A3: thứ tự thời gian mua + cấu trúc ngữ pháp của văn bản (bag-of-words)."),
    ("8. Giữ được thông tin gì?", "Toàn bộ quan hệ feature–target có ích cho dự đoán; scale/one-hot/TF-IDF đều giữ-phải-đạo-phi-tuyến."),
    ("9. Preprocessing nào có thể gây leakage?", "Impute/scale/TF-IDF fit trên toàn dữ liệu thay vì train; category revenue share (A3) sinh ra target; cho model nhìn val/test khi chọn. Tất cả đã tránh: fit trên train, chọn model trên val, test chỉ dùng 1 lần."),
    ("10. Model nào tốt nhất?", "A1: Random Forest. A2: Gradient Boosting. A3: Linear SVM (text)."),
    ("11. Vì sao chọn?", "Hiệu năng top + tiêu chí triển khai: diễn giải (RF importance), inference nhanh (GB 0.05s), text mạnh (SVM linear)."),
    ("12. Metric quan trọng nhất?", "A1: Recall (bỏ sót người bệnh). A2: MAE (sai số tiền trung bình). A3: Macro-F1 (9 lớp mất cân bằng)."),
    ("13. Persist thế nào?", "joblib 1 file chứa pipeline đầy đủ: preprocessor + model + metadata (thứ tự cột, labels, seed)."),
    ("14. Web dùng persisted model thế nào?", "Web không chứa model — gọi REST API; API load joblib, transform input đúng hệt training rồi predict."),
    ("15. Mobile giao tiếp với service thế nào?", "Mobile UI (HTML) → HTTP POST JSON → REST API → JSON response {prediction, confidence} → hiển thị kèm diễn giải."),
]
for q, a in qa:
    p(q, bold=True)
    p(a)

h2("8.2 Sáu câu riêng e-commerce")
qa2 = [
    ("1. Comment chứa thông tin gì?", "Sở thích sản phẩm thật của khách: từ khóa loại hàng khách chi tiền mua (candle/teacup/toy...)."),
    ("2. Biến thành số thế nào?", "Gộp Description theo khách thành basket_text → tokens → token IDs → trọng số TF-IDF → vector thưa 4000 chiều."),
    ("3. Token IDs nghĩa là gì?", "Chỉ số nguyên trong vocabulary học từ train — không mang giá trị số học, chỉ là danh tính của token."),
    ("4. Embedding vector nghĩa là gì?", "Mỗi chiều = một unigram/bigram; trọng số TF-IDF cao nghĩa là khách mua loại hàng đó nhiều mà nó hiếm ở các khách khác → tín hiệu sở thích mạnh."),
    ("5. Phát hiện được interest nào?", "9 nhóm: home_decor (34%), kitchen_dining, other_gifts, vintage_craft, bags_purses, toys_games, stationery, garden_outdoor."),
    ("6. Text có cải thiện so với tabular không?", "CÓ — macro-F1 tăng mạnh ở cả 3 model thử (RF 0.163→0.291; SVM 0.109→0.309; LR 0.083→0.276)."),
]
for q, a in qa2:
    p(q, bold=True)
    p(a)

# ============================ 9. REPRODUCIBILITY ============================
h1("9. Reproducibility")
table([
    ["Thành phần", "Giá trị"],
    ["Python / OS", "3.12 (Anaconda) / Windows 11"],
    ["Thư viện chính", "scikit-learn 1.9.0, pandas 2.x, numpy 1.26+, fastapi 0.141, streamlit, sentence-transformers 6.0 (chỉ mục 18.6)"],
    ["Random seed", "42 — cố định ở mọi chỗ ngẫu nhiên (split, model, sampling)"],
    ["Split", "70/15/15 stratify (A1, A3) / không stratify (A2) — random_state=42"],
    ["Dataset", "tải local trong <app>/data/ — chạy offline, kèm nguồn Kaggle ở README"],
    ["Notebook", "build từ <app>_nb_source.py bằng _build.py — cell-by-cell tái lập được"],
    ["Model files", "3 joblib pipeline (preprocess + model + metadata); deploy pin sklearn==1.9.0"],
    ["API/Web/Mobile", "mã nguồn trong <app>/api, <app>/web, <app>/mobile; deploy script deploy/deploy-all.py"],
])

# ============================ 10. CONCLUSION ============================
h1("10. Conclusion")
bullet("Bài học chính: hệ thống thông minh không phải một model đã train — mà là chuỗi "
       "Data → Clean → Represent → Learn → Evaluate → Persist → Deploy vận hành nhất quán, "
       "trong đó representation là mắt xâu nối dữ liệu thực với model.")
bullet("Thách thức kỹ thuật lớn nhất: làm sạch dữ liệu thật (0-trá hình y tế, giao dịch hủy/trùng "
       "e-commerce) và chống data leakage xuyên suốt preprocessing → split → deploy.")
bullet("Vấn đề representation quan trọng nhất: chứng minh bằng thí nghiệm rằng đổi representation "
       "đổi hiệu năng — text TF-IDF nâng macro-F1 gấp ~2 lần RF; one-hot zipcode + log-target "
       "giúp GB đạt R² 0.90.")
bullet("Bài học ML: chọn model theo metric đúng bài toán (Recall/MAE/macro-F1) và tiêu chí triển "
       "khai, không theo accuracy đơn thuần.")
bullet("Bài học deployment: cùng một pipeline đã lưu phải phục vụ cả training và serving; web và "
       "mobile cùng dùng một API — một nguồn sự thật.")
bullet("Cải tiến tương lai: (1) taxonomy category bằng model phân loại sản phẩm thay keyword "
       "rules; (2) neural embedding (MiniLM/TF-IDF hybrid) làm feature chính; (3) thu thêm dữ "
       "liệu lớp hiếm; (4) giám sát drift khi serving.")

doc.save(OUT)
print(f"Saved: {OUT}")
